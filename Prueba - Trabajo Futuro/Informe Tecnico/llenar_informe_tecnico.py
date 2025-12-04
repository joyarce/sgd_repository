# ============================================================
# llenar_informe_tecnico.py (VERSIÓN COMBINADA CORREGIDA + OPTIMIZADA)
# Genera Informe Técnico DOCX desde JSON
#
# Usa:
#   - Plantilla: Informe Tecnico.docx
#   - JSON: salida de generar_json_reportes_diarios.py
#
# Inserta:
#   - Controles globales simples (ej: nombre_proyecto, fechas, etc.)
#   - Tabla Desviaciones
#   - Tabla Actividades
#   - Curva S Final             <curva_s_final>
#   - Curva Avance Diario       <curva_avance_diario>
#   - Pareto Tiempos Perdidos   <pareto_tiempo_perdidos>
#   - Registro fotográfico      <bloque_imagen_evidencia>
#
# NOTA CLAVE:
#   La Curva S y la Curva de Avance Diario se generan SIEMPRE
#   utilizando la tabla "avance_turno" DEL ÚLTIMO REPORTE
#   (elegido por el número de reporte más alto en num_reporte_diario).
# ============================================================

import json
import argparse
import os
import tempfile
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Es necesario instalar python-docx y matplotlib (pip install python-docx matplotlib)
import matplotlib.pyplot as plt


# ------------------------------------------------------------
# NAMESPACE WORD
# ------------------------------------------------------------
def docx_ns():
    return {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


NS = docx_ns()


# ------------------------------------------------------------
# HELPER: OBTENER NÚMERO DE REPORTE COMO ENTERO
# ------------------------------------------------------------
def safe_num_reporte(doc_info):
    """
    Retorna el número de reporte como entero, a partir de
    content_controls['num_reporte_diario'].
    Si no es posible, retorna 0.
    """
    try:
        cc = doc_info.get("content_controls", {})
        nro = str(cc.get("num_reporte_diario", "0")).strip()
        return int(nro) if nro.isdigit() else 0
    except Exception:
        return 0


# ------------------------------------------------------------
# ESTABLECER VALOR DE TEXTO EN SDT
# ------------------------------------------------------------
def set_sdt_text_value(sdt, value):
    """
    Reemplaza solo el texto dentro de <w:sdtContent> del SDT.
    """
    content = sdt.find(".//w:sdtContent", NS)
    if content is None:
        return

    # Buscar solo los textos dentro de SDTContent
    textos = content.findall(".//w:t", NS)
    if not textos:
        # Si no hay <w:t>, puede ser que el SDT esté vacío o su contenido sea solo <w:p>.
        paragraphs = content.findall(".//w:p", NS)
        if paragraphs:
            para = paragraphs[0]
            para.add_run(str(value))
        return

    # Vaciar texto interno
    for t in textos:
        t.text = ""

    # Insertar valor solo en el primer t
    textos[0].text = str(value)


# ------------------------------------------------------------
# RELLENAR CONTROLES DE CONTENIDO (texto simple, globales)
# ------------------------------------------------------------
def fill_content_controls(doc, replacements):
    """
    Rellena controles de contenido de texto (SDT) cuyo tag
    exista en el diccionario 'replacements'.
    """
    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag = sdt.find(".//w:tag", NS)
        if tag is None:
            continue

        tag_val = tag.get(qn("w:val"))
        if not tag_val or tag_val not in replacements:
            continue

        new_text = str(replacements[tag_val])
        set_sdt_text_value(sdt, new_text)


# ------------------------------------------------------------
# NORMALIZAR RUTA DE IMAGEN (Windows \\ → OS)
# ------------------------------------------------------------
def normalize_image_path(path_str):
    if not path_str:
        return None

    # Limpia comillas y espacios
    path_str = str(path_str).strip().strip('"').strip("'")
    
    # Intenta con la ruta normalizada a OS (ej: / en Linux/Mac, \ en Win)
    candidate = path_str.replace("\\", os.sep)
    if os.path.exists(candidate):
        return candidate
    
    # Intenta con la ruta original si no funciona la normalizada
    if os.path.exists(path_str):
        return path_str

    return None


# ------------------------------------------------------------
# INSERTAR IMAGEN DENTRO DE UN SDT 'imagen_evidencia'
# ------------------------------------------------------------
def insert_image_into_sdt(doc, sdt, image_path, width_inches=5.5):
    """
    Inserta la imagen dentro del SDT 'imagen_evidencia'
    respetando su marco interno.
    """
    ns = NS

    # 1. Obtener contenido del SDT
    content = sdt.find(".//w:sdtContent", ns)
    if content is None:
        return

    # 2. Obtener el párrafo existente o crear uno
    paragraphs = content.findall(".//w:p", ns)
    if paragraphs:
        # Usar el primer párrafo existente dentro del contenido SDT
        p_xml = paragraphs[0]
    else:
        temp_para = doc.add_paragraph()
        p_xml = temp_para._p
        content.append(p_xml)
        
    # 3. Borrar contenido anterior del párrafo (dibujos y textos)
    for drawing in p_xml.findall(".//w:drawing", ns):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)
    for t in p_xml.findall(".//w:t", ns):
        t.text = ""

    # 4. Crear párrafo REAL para la imagen
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()

    # 5. Insertar imagen
    if image_path and os.path.exists(image_path):
        run.add_picture(image_path, width=Inches(width_inches))
    else:
        run.text = "[Imagen no encontrada]"

    # 6. Reemplazar el párrafo placeholder (p_xml) con el párrafo de la imagen
    if p_xml.getparent():
        p_xml.getparent().replace(p_xml, para._p)
    else:
        content.append(para._p)


# ------------------------------------------------------------
# OBTENER BLOQUE PLANTILLA: <bloque_imagen_evidencia>
# ------------------------------------------------------------
def get_bloque_imagen_evidencia_template(doc):
    """
    Busca el control de contenido de sección repetible que engloba
    a imagen_evidencia, num_imagen_evidencia, descripcion_imagen_evidencia, fecha_turno.
    Tag principal: 'bloque_imagen_evidencia'
    Devuelve (bloque_template, parent_element)
    """
    bloque_principal = None

    for sdt in doc.element.body.findall(".//w:sdt", NS):
        tag = sdt.find(".//w:tag", NS)
        if tag is not None and tag.get(qn("w:val")) == "bloque_imagen_evidencia":
            bloque_principal = sdt
            break

    if bloque_principal is None:
        return None, None

    parent = bloque_principal.getparent()
    bloque_template = deepcopy(bloque_principal)

    # Eliminamos el bloque original de la plantilla
    if parent is not None:
        parent.remove(bloque_principal)

    return bloque_template, parent


# ------------------------------------------------------------
# AÑADIR TÍTULO "REPORTE DIARIO N - FECHA" ANTES DEL GRUPO
# ------------------------------------------------------------
def add_reporte_title(doc, parent, num_reporte, fecha_turno):
    """
    Crea un párrafo con el título del grupo de fotos:
    "REPORTE DIARIO N - FECHA"
    y lo inserta al final del parent.
    """
    texto = f"REPORTE DIARIO N°{num_reporte}"
    if fecha_turno:
        texto += f" - {fecha_turno}"

    para = doc.add_paragraph()
    
    # Estilo (si existe); si no, usará Normal
    try:
        para.style = "Heading 3"
    except Exception:
        pass  # Usar estilo por defecto si 'Heading 3' no existe

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(texto)
    run.bold = True

    parent.append(para._p)


# ------------------------------------------------------------
# INSERTAR TODAS LAS FOTOS AGRUPADAS POR REPORTE
# ------------------------------------------------------------
def insertar_fotografias_desde_json(doc, data):
    """
    data = lista de documentos del JSON (uno por reporte diario).
    Para cada reporte:
      - Inserta un título "REPORTE DIARIO N - FECHA"
      - Inserta un bloque_imagen_evidencia por cada imagen
        del array "imagenes" de ese reporte.
    Se usa el orden de num_reporte_diario (ascendente).
    """
    bloque_template, parent = get_bloque_imagen_evidencia_template(doc)
    if bloque_template is None or parent is None:
        print("No se encontró el bloque 'bloque_imagen_evidencia' en la plantilla.")
        return

    # Ordenar los reportes por número de reporte
    data_ordenada = sorted(data, key=safe_num_reporte)

    for doc_info in data_ordenada:
        cc = doc_info.get("content_controls", {})
        imagenes = doc_info.get("imagenes", []) or []

        if not imagenes:
            continue

        num_reporte = cc.get("num_reporte_diario", "?")
        fecha_turno_doc = cc.get("fecha_turno", "")

        # Título del grupo
        add_reporte_title(doc, parent, num_reporte, fecha_turno_doc)

        # Bloque por cada imagen
        for idx, img_info in enumerate(imagenes, start=1):
            new_block = deepcopy(bloque_template)
            parent.append(new_block)

            img_path_raw = img_info.get("img", "")
            desc = img_info.get("descripcion", "") or ""
            # Si la foto tiene fecha, usarla; si no, usar la fecha del reporte
            fecha_img = img_info.get("fecha") or fecha_turno_doc or "" 

            img_path = normalize_image_path(img_path_raw)

            for sdt in new_block.findall(".//w:sdt", NS):
                tag_elem = sdt.find(".//w:tag", NS)
                if tag_elem is None:
                    continue

                tag_val = tag_elem.get(qn("w:val"))

                if tag_val == "num_imagen_evidencia":
                    set_sdt_text_value(sdt, f"FOTO {idx}")
                elif tag_val == "descripcion_imagen_evidencia":
                    set_sdt_text_value(sdt, desc)
                elif tag_val == "fecha_turno":
                    set_sdt_text_value(sdt, fecha_img)
                elif tag_val == "imagen_evidencia":
                    insert_image_into_sdt(doc, sdt, img_path, width_inches=5.5)


# ------------------------------------------------------------
# INSERTAR IMAGEN EN CONTROL SIMPLE (GRÁFICOS)
# ------------------------------------------------------------
def replace_sdt_with_image(doc, tag_name, image_path, width_inches=6):
    """
    Inserta una imagen dentro del SDT existente sin crear nuevos controles.
    Funciona para curva_s_final, curva_avance_diario y pareto_tiempo_perdidos.
    """
    from docx.shared import Inches as DocxInches
    from docx.oxml import OxmlElement
    ns = NS

    if not image_path or not os.path.exists(image_path):
        print(f"[WARN] Imagen no encontrada para '{tag_name}': {image_path}")
        return

    # Buscar el SDT correcto
    for sdt in doc.element.body.iter(qn("w:sdt")):
        tag_el = sdt.find(".//w:tag", ns)
        if tag_el is None or tag_el.get(qn("w:val")) != tag_name:
            continue

        content = sdt.find(".//w:sdtContent", ns)
        if content is None:
            continue

        # 1. ELIMINAR dibujos y textos previos
        for drawing in content.findall(".//w:drawing", ns):
            parent = drawing.getparent()
            if parent is not None:
                parent.remove(drawing)

        for t in content.findall(".//w:t", ns):
            t.text = ""

        # 2. Tomar o crear el párrafo dentro del SDT
        paragraphs = content.findall("./w:p", ns)
        if paragraphs:
            p = paragraphs[0]
        else:
            p = OxmlElement("w:p")
            content.append(p)

        # 3. Crear un 'w:r' real dentro del mismo párrafo
        r = OxmlElement("w:r")
        p.append(r)

        # 4. Generar temporalmente la imagen con python-docx
        #    (solo para obtener el w:drawing)
        temp_para = doc.add_paragraph()
        temp_r = temp_para.add_run()
        temp_r.add_picture(image_path, width=DocxInches(width_inches))

        # 5. Extraer el w:drawing generado
        drawing_el = temp_para._p.xpath(".//w:drawing")
        if drawing_el:
            r.append(drawing_el[0])

        # 6. Eliminar el párrafo temporal creado
        temp_para._p.getparent().remove(temp_para._p)

        return  # terminamos, solo 1 SDT


# ------------------------------------------------------------
# AÑADIR FILA A TABLA WORD
# ------------------------------------------------------------
def append_table_row(table, values):
    row = table.add_row()
    # Asegurar que el número de celdas coincida con el número de valores
    if len(row.cells) != len(values):
        # Advertencia: La tabla no tiene el mismo número de columnas que los datos a insertar
        pass 
        
    for i, v in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = "" if v is None else str(v)


# ------------------------------------------------------------
# BUSCAR TABLA POR ENCABEZADO
# ------------------------------------------------------------
def find_table_by_header(doc, header_row_text):
    """
    Busca una tabla cuya primera fila (encabezado) coincida
    con el texto de referencia, ignorando espacios y mayúsculas.
    El texto de referencia debe usar '|' como separador de celdas.
    """
    header_clean = header_row_text.replace(" ", "").lower()

    for table in doc.tables:
        if not table.rows:
            continue
            
        first = table.rows[0].cells
        # Unir el texto de las celdas de la primera fila con '|'
        join = "|".join(c.text.replace(" ", "").lower() for c in first)
        
        if join == header_clean:
            return table
    return None


# ------------------------------------------------------------
# MAIN
# ------------------------------------------------------------
def main(json_file, template_path, output_path):
    # ========================
    # Cargar JSON
    # ========================
    print("\n=== Cargando JSON ===")
    try:
        with open(json_file, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        print(f"ERROR: No se pudo cargar el archivo JSON {json_file}. Error: {e}")
        return

    if not data:
        print("ADVERTENCIA: El JSON está vacío. No se generará el informe.")
        return
        
    print(f"Turnos cargados: {len(data)}")

    # ========================
    # Ordenar por N° de reporte
    # ========================
    data_ordenada = sorted(data, key=safe_num_reporte)
    primer = data_ordenada[0].get("content_controls", {})
    ultimo = data_ordenada[-1].get("content_controls", {})

    # ========================
    # Cargar plantilla
    # ========================
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"ERROR: No se pudo cargar la plantilla DOCX {template_path}. Error: {e}")
        return

    # ========================
    # 1. Relleno CC globales
    # ========================
    replacements = {
        "nombre_proyecto": primer.get("nombre_proyecto", ""),
        "fecha_inicio_ejecucion": primer.get("fecha_turno", ""),
        "fecha_cierre_proyecto": ultimo.get("fecha_turno", "")
    }
    replacements = {k: v for k, v in replacements.items() if v}
    if replacements:
        print("Rellenando controles de contenido globales...")
        fill_content_controls(doc, replacements)

    # ========================
    # 2. Tabla Desviaciones
    # ========================
    header_desv = "ID|Fecha|Inicio|Fin|Hr|R.Crítica|Resp.|Tipo|Descripción"
    tbl_desv = find_table_by_header(doc, header_desv)

    if tbl_desv is None:
        print("ADVERTENCIA: No se encontró tabla de desviaciones.")
    else:
        print("Rellenando tabla de desviaciones...")
        # Eliminar las filas de datos de ejemplo (manteniendo el encabezado)
        for i in range(len(tbl_desv.rows) - 1, 0, -1):
            tbl_desv._element.remove(tbl_desv.rows[i]._element)

        id_ctr = 1
        for rep in data_ordenada:
            fecha = rep.get("content_controls", {}).get("fecha_turno", "")
            desv_rows = rep.get("excel_tables", {}).get("desviaciones_turno", {}).get("rows", [])

            for row in desv_rows:
                append_table_row(
                    tbl_desv,
                    [
                        id_ctr,
                        fecha,
                        row.get("H. INICIO", ""),
                        row.get("H. FIN", ""),
                        row.get("H. TOTAL", ""),
                        row.get("R. CRIT", row.get("R. CRÍT", "")),  # Aceptar R. CRIT o R. CRÍT
                        "",  # Columna Resp.
                        "",  # Columna Tipo
                        row.get("DESCRIPCION", row.get("DESCRIPCIÓN", ""))  # Aceptar DESCRIPCION o DESCRIPCIÓN
                    ]
                )
                id_ctr += 1

    # ========================
    # 3. Tabla Actividades
    # ========================
    header_acts = "Fecha turno|Jornada|Actividad"
    tbl_acts = find_table_by_header(doc, header_acts)

    if tbl_acts is None:
        print("ADVERTENCIA: No se encontró tabla actividades.")
    else:
        print("Rellenando tabla de actividades...")
        # Eliminar las filas de datos de ejemplo (manteniendo el encabezado)
        for i in range(len(tbl_acts.rows) - 1, 0, -1):
            tbl_acts._element.remove(tbl_acts.rows[i]._element)

        for rep in data_ordenada:
            cc = rep.get("content_controls", {})
            fecha = cc.get("fecha_turno", "")
            jornada = cc.get("tipo_jornada", "")
            acts = rep.get("excel_tables", {}).get("act_realizadas", {}).get("rows", [])

            for a in acts:
                append_table_row(
                    tbl_acts,
                    [fecha, jornada, a.get("DESCRIPCION", a.get("DESCRIPCIÓN", ""))]
                )

    # ========================
    # 4. GRÁFICOS
    # ========================
    print("\nGenerando gráficos de avance...")

    def to_float(v):
        if v is None:
            return 0.0
        if isinstance(v, (int, float)):
            return float(v)
        s = str(v).replace("%", "").replace(",", ".").strip()
        try:
            return float(s)
        except Exception:
            return 0.0

    # ---- Avance Turno (Curva S y curva diaria) ----
    # ---- Avance Turno (Curva S y curva diaria) ----
    print("\nGenerando gráficos de avance...")

    ultimo_reporte = data_ordenada[-1]
    avance = ultimo_reporte.get("excel_tables", {}).get("avance_turno", {})
    filas = avance.get("rows", [])

    tmp_files = []

    if filas:
        filas_validas = []
        for idx, row in enumerate(filas, 1):
            nro = row.get("N° REPORTE DIARIO", row.get("N°REPORTEDIARIO", idx))
            nro_str = str(nro).strip()
            nro_i = int(nro_str) if nro_str.isdigit() else idx
            filas_validas.append((nro_i, row))

        filas_validas.sort(key=lambda x: x[0])

        xs = []
        plan_ac = []
        real_ac = []
        plan_d = []
        real_d = []

        for n, r in filas_validas:
            xs.append(n)

            # --- Curva S (acumulado) ---
            pa_raw = to_float(r.get("AVANCE PLANIFICADO (ACUMULADO)", 0))
            ra_raw = to_float(r.get("AVANCE REAL (ACUMULADO)", 0))

            # Corrige valores 0–1 → porcentaje (0–100)
            pa = pa_raw * 100 if 0 <= pa_raw <= 1 else pa_raw
            ra = ra_raw * 100 if 0 <= ra_raw <= 1 else ra_raw

            plan_ac.append(pa)
            real_ac.append(ra)

            # --- Curva diaria ---
            pd_raw = to_float(r.get("% AVANCE DIARIO PLANIFICADO",
                                    r.get("%AVANCEDIARIOPLANIFICADO", 0)))
            rd_raw = to_float(r.get("% AVANCE DIARIO REAL",
                                    r.get("%AVANCEDIARIOREAL", 0)))

            pd = pd_raw * 100 if 0 <= pd_raw <= 1 else pd_raw
            rd = rd_raw * 100 if 0 <= rd_raw <= 1 else rd_raw

            plan_d.append(pd)
            real_d.append(rd)

        # --- Generar Curva S ---
        plt.figure(figsize=(8, 5))
        plt.plot(xs, plan_ac, marker="o", linewidth=2, label="Plan Acumulado (%)")
        plt.plot(xs, real_ac, marker="o", linewidth=2, label="Real Acumulado (%)")
        plt.title("Curva S - Avance Acumulado")
        plt.xlabel("N° Reporte Diario")
        plt.ylabel("Avance Acumulado (%)")
        plt.grid(True, linestyle="--", alpha=0.6)
        plt.legend()

        tmp_s = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        plt.savefig(tmp_s.name, dpi=160, bbox_inches="tight")
        plt.close()
        tmp_files.append(tmp_s.name)
        replace_sdt_with_image(doc, "curva_s_final", tmp_s.name)

        # --- Generar Curva Avance Diario ---
        plt.figure(figsize=(8, 5))
        plt.plot(xs, plan_d, marker="o", linewidth=2, label="Plan Diario (%)")
        plt.plot(xs, real_d, marker="o", linewidth=2, label="Real Diario (%)")
        plt.title("Curva Avance Diario")
        plt.xlabel("N° Reporte Diario")
        plt.ylabel("Avance Diario (%)")
        plt.grid(True, linestyle="--", alpha=0.6)
        plt.legend()

        tmp_d = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        plt.savefig(tmp_d.name, dpi=160, bbox_inches="tight")
        plt.close()
        tmp_files.append(tmp_d.name)
        replace_sdt_with_image(doc, "curva_avance_diario", tmp_d.name)

    else:
        print("ADVERTENCIA: El último reporte no tiene filas en 'avance_turno'.")


    # ========================
    # 5. PARETO TIEMPOS PERDIDOS
    # ========================
    print("Generando Pareto...")

    def time_to_hours(t):
        if not t or t in ("", "None"):
            return 0
        try:
            # Asume formato HH:MM:SS
            parts = str(t).split(":")
            if len(parts) == 3:
                h, m, s = parts
                return int(h) + int(m) / 60 + int(s) / 3600
            elif len(parts) == 2:
                # Asume formato MM:SS
                m, s = parts
                return int(m) / 60 + int(s) / 3600
            else:
                return 0
        except Exception:
            return 0

    pareto = {}

    for rep in data_ordenada:
        desvs = rep.get("excel_tables", {}).get("desviaciones_turno", {}).get("rows", [])
        for r in desvs:
            desc = r.get("DESCRIPCION", r.get("DESCRIPCIÓN", "Sin descripción"))
            horas = time_to_hours(r.get("H. TOTAL", "0:00:00"))
            pareto[desc] = pareto.get(desc, 0) + horas

    if pareto and sum(pareto.values()) > 0:
        items = sorted(pareto.items(), key=lambda x: x[1], reverse=True)
        labels = [i[0] for i in items]
        vals = [i[1] for i in items]

        total = sum(vals)
        acum = []
        run = 0
        for v in vals:
            run += v
            acum.append(run / total * 100 if total > 0 else 0)

        fig, ax1 = plt.subplots(figsize=(10, 6))

        # Barras (Horas)
        ax1.bar(labels, vals, alpha=0.7)
        ax1.set_xlabel("Descripción de Tiempo Perdido")
        ax1.set_ylabel("Horas Perdidas (H. TOTAL)")
        ax1.tick_params(axis='y')

        # Rotación y alineación de etiquetas del eje X
        plt.setp(ax1.get_xticklabels(), rotation=45, ha="right", rotation_mode="anchor")

        # Línea (Acumulado %)
        ax2 = ax1.twinx()
        ax2.plot(labels, acum, marker="o", linewidth=2, label="Acumulado (%)")
        ax2.set_ylabel("Porcentaje Acumulado (%)")
        ax2.tick_params(axis='y')
        ax2.set_ylim(0, 100)
        ax2.grid(True, linestyle="--", alpha=0.5)

        plt.title("Pareto de Tiempos Perdidos")

        tmp_p = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        plt.tight_layout()
        plt.savefig(tmp_p.name, dpi=160)
        plt.close(fig)
        tmp_files.append(tmp_p.name)
        replace_sdt_with_image(doc, "pareto_tiempo_perdidos", tmp_p.name)
    else:
        print("ADVERTENCIA: No hay datos válidos para generar el Pareto de Tiempos Perdidos.")

    # ========================
    # 6. REGISTRO FOTOGRÁFICO
    # ========================
    print("\nInsertando registro fotográfico...")
    insertar_fotografias_desde_json(doc, data_ordenada)

    # ========================
    # GUARDAR
    # ========================
    try:
        doc.save(output_path)
        print("\n✔ INFORME TÉCNICO generado y guardado:")
        print(output_path)
    except Exception as e:
        print(f"ERROR: No se pudo guardar el documento en {output_path}. "
              f"Asegúrese de que no esté abierto. Error: {e}")
    
    # ========================
    # LIMPIEZA
    # ========================
    for tmp in tmp_files:
        try:
            os.remove(tmp)
        except Exception:
            pass  # Ignorar errores de borrado de temporales


# ------------------------------------------------------------
# CLI
# ------------------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generar Informe Técnico DOCX desde el JSON de reportes diarios."
    )
    parser.add_argument("--json", required=True, help="Ruta al JSON de entrada.")
    parser.add_argument("--template", required=True, help="Ruta a la plantilla DOCX.")
    parser.add_argument(
        "--out",
        default="INFORME_TECNICO_COMPLETO.docx",
        help="Ruta de salida del Informe Técnico."
    )
    args = parser.parse_args()
    main(args.json, args.template, args.out)
