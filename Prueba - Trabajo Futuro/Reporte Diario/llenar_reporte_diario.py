# ========================================================================
# llenar_reporte_diario_desde_json.py — VERSIÓN UNIVERSAL FINAL
# Consume estructura_reporte.json y rellena el DOCX original.
# Reemplaza TODOS los SDT (controles), rellena excels embebidos,
# mantiene estilos, fórmulas, imágenes, nombres definidos,
# genera datos sintéticos de turno y curva S.
# ========================================================================

import os
import json
import zipfile
import random
import math
from datetime import datetime, timedelta
from io import BytesIO
from copy import copy, deepcopy
import xml.etree.ElementTree as ET

from openpyxl import load_workbook
from openpyxl.utils import range_boundaries, get_column_letter
from openpyxl.utils.exceptions import InvalidFileException

from tempfile import NamedTemporaryFile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml.ns import qn as qn_docx
from docx.shared import Inches

# Namespace Word
NS_W = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


# ------------------------------------------------------------------------
# CATÁLOGOS DE DATOS SINTÉTICOS (MINERÍA / OVERHAUL)
# ------------------------------------------------------------------------
PROYECTOS = [
    "Overhaul Chancador Primario MP1000",
    "Reparación Molino SAG 40x22",
    "Cambio Coraza Molino Bolas 26x38",
    "Mantención Harnero Banana DF501",
    "Reemplazo Sistema Hidráulico HP300",
    "Major Overhaul Chancador GP500",
]

REFERENCIAS_UBICACIONES = [
    ("CH-001", "CHANCADORA PRIMARIA - ZONA NORTE"),
    ("CH-002", "CHANCADORA SECUNDARIA - ZONA SUR"),
    ("SG-001", "MOLINO SAG - ZONA ESTE"),
    ("MB-001", "MOLINO BOLAS - ZONA OESTE"),
    ("HZ-010", "HARNEO BANANA - SECTOR INTERMEDIO"),
]

TURNOS_POSSIBLES = [
    ("07:00", "19:00"),
    ("08:00", "20:00"),
    ("06:30", "18:30"),
    ("21:00", "09:00"),
    ("22:00", "10:00"),
    ("20:30", "08:30"),
]

REPRESENTANTES_METSO = [
    ("Juan Soto", "Administrador de Contrato"),
    ("Camila Rojas", "Supervisor Mecánico"),
    ("Pedro González", "Ingeniero de Proyecto"),
    ("Ana Villalobos", "Jefa HSE"),
]

REPRESENTANTES_CLIENTE = [
    ("Luis Arrieta", "Jefe Mantención Planta"),
    ("Francisca Rivas", "Superintendente Mantención"),
    ("Rodrigo Olea", "Jefe Turno Cliente"),
    ("Natalia Fuentes", "Planificador Cliente"),
]


# =====================================================================
# AUXILIAR: CALCULAR HORAS DE TURNO
# =====================================================================
def _calculate_shift_hours(start_time_str, end_time_str):
    """
    Calcula la duración del turno en horas decimales (hh:mm -> float).
    Maneja turnos que cruzan la medianoche.
    Ejemplo: 21:00 a 09:00 = 12 horas.
    """
    start_time = datetime.strptime(start_time_str, "%H:%M")
    end_time = datetime.strptime(end_time_str, "%H:%M")

    if end_time < start_time:
        end_time += timedelta(days=1)

    duration = end_time - start_time
    total_hours = duration.total_seconds() / 3600.0
    return total_hours


# =====================================================================
# ACTUALIZAR NOMBRES DEFINIDOS EN XLSX (HH_TURNO)
# =====================================================================
def update_embedded_xlsx_names(xlsx_bytes, name_values):
    """
    Actualiza los valores de los nombres definidos de Excel.
    name_values es un diccionario {nombre: valor}
    """
    try:
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=False)
    except InvalidFileException:
        return xlsx_bytes

    modified = False

    for name, name_obj in wb.defined_names.items():
        if name in name_values:
            try:
                dest_sheet, dest_coord = next(name_obj.destinations)
            except StopIteration:
                continue

            ws = wb[dest_sheet]
            cell_ref = dest_coord
            cell = ws[cell_ref]
            cell.value = name_values[name]
            cell.number_format = "[h]:mm"
            modified = True

    if modified:
        out = BytesIO()
        wb.save(out)
        return out.getvalue()

    return xlsx_bytes


# =====================================================================
# AJUSTE DE FÓRMULAS (cuando se insertan filas en tablas)
# =====================================================================
def _shift_formula_rows(formula, src_row, dst_row):
    """
    Ajusta referencias A1 en fórmulas.
    Las referencias estructuradas de tabla ([[#This Row],...]) NO deben tocarse.
    """
    import re

    if "[#This Row]" in formula or "[[" in formula:
        return formula

    row_offset = dst_row - src_row
    pattern = re.compile(r"(\$?[A-Z]{1,3})(\d+)")

    def repl(match):
        col = match.group(1)
        row = int(match.group(2))
        new_row = row + row_offset
        return f"{col}{new_row}"

    return pattern.sub(repl, formula)


# =====================================================================
# UTILIDADES DE TABLA
# =====================================================================
def _ref_to_bounds(ref):
    return range_boundaries(ref)


def _get_table_headers(ws, ref):
    min_col, min_row, max_col, max_row = _ref_to_bounds(ref)
    header_row = min_row
    headers = {}

    for col in range(min_col, max_col + 1):
        cell = ws.cell(row=header_row, column=col)
        if cell.value:
            headers[str(cell.value).strip()] = col

    return headers, min_row, max_row


def _is_hour_format_valid(value):
    try:
        f_val = float(value)
        return f_val >= 0 and f_val < 1
    except (ValueError, TypeError):
        return False


def _is_integer_valid(value):
    try:
        i_val = int(value)
        return i_val >= 0
    except (ValueError, TypeError):
        return False


# =====================================================================
# EXPANDIR TABLA EXCEL AUTOMÁTICAMENTE
# =====================================================================
def ensure_table_min_rows(ws, tbl, required_rows):
    """
    Garantiza que la tabla Excel tenga al menos 'required_rows' filas de datos.
    Si no, inserta filas, actualiza tbl.ref y copia formato/fórmulas
    de la PRIMERA FILA DE DATOS de la tabla.
    """
    min_col, min_row, max_col, max_row = range_boundaries(tbl.ref)
    header_row = min_row
    first_data_row = header_row + 1

    current_rows = max_row - header_row
    missing = required_rows - current_rows
    if missing <= 0:
        return header_row, max_row

    insert_at = max_row + 1
    ws.insert_rows(insert_at, missing)
    new_max_row = max_row + missing

    for col in range(min_col, max_col + 1):
        src_cell = ws.cell(row=first_data_row, column=col)
        for new_row in range(insert_at, new_max_row + 1):
            dst_cell = ws.cell(row=new_row, column=col)
            dst_cell._style = src_cell._style

            src_val = src_cell.value
            if isinstance(src_val, str) and src_val.startswith("="):
                if "[#This Row]" in src_val or "[[" in src_val:
                    dst_cell.value = src_val
                else:
                    dst_cell.value = _shift_formula_rows(
                        src_val, first_data_row, new_row
                    )
            else:
                dst_cell.value = None

    new_ref = (
        f"{get_column_letter(min_col)}{min_row}:"
        f"{get_column_letter(max_col)}{new_max_row}"
    )
    tbl.ref = new_ref
    return header_row, new_max_row


# =====================================================================
# 1) RELLENAR CONTROLES DE CONTENIDO (document.xml) POR TAG
# =====================================================================
def fill_content_controls(document_xml_bytes, context):
    root = ET.fromstring(document_xml_bytes)

    for sdt in root.findall(".//w:sdt", NS_W):
        tag_elem = sdt.find("./w:sdtPr/w:tag", NS_W)
        if tag_elem is None:
            continue

        tag_val = tag_elem.get(f"{{{NS_W['w']}}}val")

        # Caso especial: plantilla nombre_representante_reporte_diario
        if (
            tag_val == "nombre_representante_reporte_diario"
            and "nombre_representante_metso_reporte_diario" in context
        ):
            tag_val = "nombre_representante_metso_reporte_diario"

        if not tag_val or tag_val not in context:
            continue

        new_text = str(context[tag_val])
        texts = sdt.findall(".//w:sdtContent//w:t", NS_W)
        if not texts:
            continue

        for t in texts:
            t.text = ""
        texts[0].text = new_text

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


# =====================================================================
# 2) ACTUALIZAR TABLA 'avance_turno' (curva S)
# =====================================================================
def update_avance_turno_xlsx(xlsx_bytes, doc_index, total_docs, daily_plan, daily_real):
    # NOMBRES EXACTOS IGUALES AL ARCHIVO EXCEL (SIN ESPACIOS RAROS)
    col_name_plan = "% AVANCE DIARIO PLANIFICADO"
    col_name_real = "% AVANCE DIARIO REAL"  # <-- CORREGIDO, SIN ESPACIO FINAL
    col_name_acum_plan = "AVANCE PLANIFICADO (ACUMULADO)"
    col_name_acum_real = "AVANCE REAL (ACUMULADO)"
    col_name_n_reporte = "N° REPORTE DIARIO"

    try:
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=False)
    except InvalidFileException:
        return xlsx_bytes

    found = False

    for ws in wb.worksheets:
        for tbl in ws._tables.values():
            if tbl.name != "avance_turno":
                continue

            found = True
            # Asegurar que hay tantas filas como N reportes
            header_row, max_row = ensure_table_min_rows(ws, tbl, total_docs)

            ref = tbl.ref
            headers, header_row, max_row = _get_table_headers(ws, ref)
            total_rows = max_row - header_row

            col_map = {
                col_name_n_reporte: headers.get(col_name_n_reporte),
                col_name_plan: headers.get(col_name_plan),
                col_name_real: headers.get(col_name_real),
                col_name_acum_plan: headers.get(col_name_acum_plan),
                col_name_acum_real: headers.get(col_name_acum_real),
            }

            # Si falta alguna columna clave, no tocamos nada
            if not all(col_map.values()):
                return xlsx_bytes

            for i in range(1, total_rows + 1):
                excel_row = header_row + i

                if i <= doc_index:
                    p = float(daily_plan[i - 1])
                    r = float(daily_real[i - 1])
                    acum_p = sum(daily_plan[:i])
                    acum_r = sum(daily_real[:i])

                    ws.cell(excel_row, col_map[col_name_n_reporte]).value = i
                    ws.cell(excel_row, col_map[col_name_plan]).value = p
                    ws.cell(excel_row, col_map[col_name_real]).value = r

                    if i == total_docs:
                        # Último reporte: acumulado = 100% (1.0)
                        ws.cell(excel_row, col_map[col_name_acum_plan]).value = 1.0
                        ws.cell(excel_row, col_map[col_name_acum_real]).value = 1.0
                    else:
                        ws.cell(excel_row, col_map[col_name_acum_plan]).value = acum_p
                        ws.cell(excel_row, col_map[col_name_acum_real]).value = acum_r

                elif i <= total_docs:
                    # Filas intermedias > doc_index pero <= total_docs: limpiar
                    for col in col_map.values():
                        ws.cell(excel_row, col).value = None
                else:
                    # Filas sobrantes fuera de la tabla lógica
                    for col in col_map.values():
                        ws.cell(excel_row, col).value = None

    if not found:
        return xlsx_bytes

    out = BytesIO()
    wb.save(out)
    return out.getvalue()



# =====================================================================
# 3) ACTUALIZAR OTRAS TABLAS EMBEBIDAS
# =====================================================================
def update_data_rows(ws, table_name, data_rows, header_row, max_row, headers):
    HOUR_COLS = ["H.M. ASIGN.", "H. M. OPER.", "H.M. S-BY", "H. INICIO", "H. FIN"]
    INT_COLS = ["CANT", "CANT. PLAN.", "CANT. REAL"]

    for idx, row_data in enumerate(data_rows):
        excel_row = header_row + idx + 1
        if excel_row > max_row:
            continue

        for col_name, value in row_data.items():
            col_name_strip = str(col_name).strip()
            if col_name_strip not in headers:
                continue

            col = headers[col_name_strip]

            if col_name_strip in HOUR_COLS:
                if not _is_hour_format_valid(value):
                    raise ValueError(
                        f"Tabla '{table_name}': '{col_name_strip}' fila {idx+1} debe "
                        f"ser float >=0 <1. Valor: {value}"
                    )

            if col_name_strip in INT_COLS:
                if not _is_integer_valid(value):
                    raise ValueError(
                        f"Tabla '{table_name}': '{col_name_strip}' fila {idx+1} debe "
                        f"ser entero >=0. Valor: {value}"
                    )

            if col_name_strip in [
                "HH PLAN.",
                "HH REAL",
                "H. TOTAL",
                "DISPONIBILIDAD",
                "UTILIZACIÓN",
                "USO EFECTIVO",
                "STAND-BY",
            ]:
                continue

            ws.cell(excel_row, col).value = value

    rows_to_clear_start = header_row + len(data_rows) + 1
    for row in range(rows_to_clear_start, max_row + 1):
        for col in headers.values():
            ws.cell(row=row, column=col).value = None


def update_embedded_xlsx_tables(xlsx_bytes, tables_context):
    try:
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=False)
    except InvalidFileException:
        return xlsx_bytes

    modified = False

    for ws in wb.worksheets:
        for tbl in ws._tables.values():
            name = tbl.name
            if name not in tables_context:
                continue
            if name == "avance_turno":
                continue

            modified = True
            rows = tables_context[name]
            required_rows = len(rows)

            try:
                if required_rows > 0:
                    ensure_table_min_rows(ws, tbl, required_rows)

                ref = tbl.ref
                headers, header_row, max_row = _get_table_headers(ws, ref)
                update_data_rows(ws, name, rows, header_row, max_row, headers)
            except ValueError as e:
                raise e
            except Exception as e:
                print(f"Error procesando tabla '{name}': {e}")

    if modified:
        out = BytesIO()
        wb.save(out)
        return out.getvalue()

    return xlsx_bytes


# =====================================================================
# 4A) GENERAR IMAGEN DE EVIDENCIA
# =====================================================================
def _crear_imagen_evidencia(texto):
    """
    Genera una imagen PNG de fondo blanco con texto centrado en negro.
    Devuelve la ruta del archivo temporal.
    """
    width, height = 1200, 800
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("arial.ttf", 48)
    except Exception:
        font = ImageFont.load_default()

    lines = texto.split("\n")
    line_sizes = []

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        line_sizes.append((w, h))

    total_h = sum(h for _, h in line_sizes) + (len(lines) - 1) * 10
    y = (height - total_h) // 2

    for (w, h), line in zip(line_sizes, lines):
        x = (width - w) // 2
        draw.text((x, y), line, fill="black", font=font)
        y += h + 10

    tmp = NamedTemporaryFile(delete=False, suffix=".png")
    img.save(tmp.name, "PNG")
    return tmp.name


# =====================================================================
# 4B) INSERTAR MÚLTIPLES EVIDENCIAS EN EL DOCX
# =====================================================================
def _insertar_evidencias_en_docx(doc_path, num_reporte_diario, total_evidencias=4):
    """
    Inserta múltiples evidencias fotográficas en el documento.
    Duplica el bloque SDT que contiene:
        - num_imagen_evidencia
        - imagen_evidencia
        - descripcion_imagen_evidencia
    """
    doc = Document(doc_path)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    bloque_principal = None
    for sdt in doc.element.body.findall(".//w:sdt", ns):
        tags = [t.get(qn_docx("w:val")) for t in sdt.findall(".//w:tag", ns)]
        if "bloque_imagen_evidencia" in tags:
            bloque_principal = sdt
            break

    if bloque_principal is None:
        print("❌ No se encontró el bloque SDT de evidencia con tag 'bloque_imagen_evidencia'.")
        return

    bloque_template = deepcopy(bloque_principal)
    parent = bloque_principal.getparent()
    parent.remove(bloque_principal)

    for i in range(1, total_evidencias + 1):
        new_block = deepcopy(bloque_template)
        parent.append(new_block)

        titulo_imagen = (
            f"Foto Evidencia {i} de {total_evidencias}\n"
            f"Reporte Turno {num_reporte_diario}"
        )
        desc_imagen = (
            f"Registro fotográfico {i} de {total_evidencias} "
            f"correspondiente al Reporte Diario {num_reporte_diario}."
        )

        img_path = _crear_imagen_evidencia(titulo_imagen)

        for sdt in new_block.findall(".//w:sdt", ns):
            tag = sdt.find(".//w:tag", ns)
            if tag is None:
                continue

            tag_val = tag.get(qn_docx("w:val"))

            if tag_val == "num_imagen_evidencia":
                for t in sdt.findall(".//w:t", ns):
                    t.text = ""
                t = sdt.find(".//w:t", ns)
                if t is not None:
                    t.text = str(i)

            if tag_val == "descripcion_imagen_evidencia":
                for t in sdt.findall(".//w:t", ns):
                    t.text = ""
                t = sdt.find(".//w:t", ns)
                if t is not None:
                    t.text = desc_imagen

            if tag_val == "imagen_evidencia":
                for d in sdt.findall(".//w:drawing", ns):
                    d.getparent().remove(d)
                for t in sdt.findall(".//w:t", ns):
                    t.text = ""

                content = sdt.find(".//w:sdtContent", ns)
                if content is not None:
                    paragraphs = content.findall(".//w:p", ns)
                    if paragraphs:
                        p = paragraphs[0]
                        para_obj = doc._body.add_paragraph()
                        para_obj.alignment = 1
                        run = para_obj.add_run()
                        run.add_picture(img_path, width=Inches(5.5))
                        p.addnext(para_obj._p)
                    else:
                        para_obj = doc._body.add_paragraph()
                        para_obj.alignment = 1
                        run = para_obj.add_run()
                        run.add_picture(img_path, width=Inches(5.5))
                        content.append(para_obj._p)

    doc.save(doc_path)


# =====================================================================
# 4C) CURVA S CLÁSICA (para daily_plan / daily_real)
# =====================================================================
def generar_curvas_s_clasica(N, ruido_real=0.10):
    """
    Genera daily_plan y daily_real con forma de curva S (sigmoidal).
    - daily_plan: suave, simétrica, suma 1.0
    - daily_real: daily_plan + ruido, normalizado a 1.0
    """
    k = N / 2
    a = N / 10

    S = [1 / (1 + math.exp(-(i - k) / a)) for i in range(N + 1)]
    S_min = S[0]
    S_max = S[-1]
    S = [(x - S_min) / (S_max - S_min) for x in S]

    daily_plan = [S[i] - S[i - 1] for i in range(1, N + 1)]
    total = sum(daily_plan)
    daily_plan = [x / total for x in daily_plan]

    daily_real = []
    for p in daily_plan:
        delta = random.uniform(-ruido_real, ruido_real)
        daily_real.append(max(0.0, p + delta))

    total_r = sum(daily_real)
    daily_real = [r / total_r for r in daily_real]

    return daily_plan, daily_real


# =====================================================================
# 5) GENERADORES DE TABLAS (datos sintéticos coherentes)
# =====================================================================
def generar_personal_directo():
    base = [
        ("Supervisor Mecánico", 1),
        ("Téc. Senior Mecánico", 2),
        ("Téc. Mecánico", 4),
        ("Soldador Especialista", 2),
        ("Ayudante Mecánico", 3),
        ("Téc. Lubricación", 1),
        ("Téc. Hidráulico", 1),
    ]
    personal = []
    for cargo, base_cant in base:
        variacion = random.choice([-1, 0, 1])
        cant_plan = max(1, base_cant + variacion)
        cant_real = max(0, cant_plan + random.choice([-1, 0, 1]))
        personal.append(
            {
                "CANT": cant_plan,
                "CARGO": cargo,
                "CANT. PLAN.": cant_plan,
                "CANT. REAL": cant_real,
            }
        )
    return personal


def generar_personal_indirecto():
    cargos = [
        "Prevencionista HSE",
        "Bodeguero / Logística",
        "Planificador",
        "Ingeniero de Proyecto",
        "Administrador de Contrato",
    ]
    personal = []
    for cargo in cargos:
        cant_plan = 1
        cant_real = 1 if random.random() > 0.2 else 0
        personal.append(
            {
                "CANT": cant_plan,
                "CARGO": cargo,
                "CANT. PLAN.": cant_plan,
                "CANT. REAL": cant_real,
            }
        )
    return personal


def generar_equipos_auxiliares():
    equipos = [
        ("Camión Pluma", 0.45, 0.50),
        ("Camioneta", 0.50, 0.50),
        ("Grúa Horquilla", 0.40, 0.48),
        ("Plataforma Elevadora", 0.35, 0.45),
        ("Compresor", 0.50, 0.50),
        ("Torre de Iluminación", 0.45, 0.50),
    ]
    motivos = [
        "Condición de terreno",
        "Esperando instrucciones",
        "Clima adverso",
        "Interferencia de otros equipos",
        "Falta de operador",
    ]
    lista = []
    for nombre, min_op, max_op in equipos:
        asign = 0.50
        oper = round(random.uniform(min_op, max_op), 3)
        standby = round(max(0.0, asign - oper), 3)
        motivo = random.choice(motivos) if standby > 0 else ""
        lista.append(
            {
                "CANT": random.choice([1, 1, 2]),
                "DESCRIPCIÓN": nombre,
                "H.M. ASIGN.": asign,
                "H. M. OPER.": oper,
                "H.M. S-BY": standby,
                "OBS. / MOTIVO STAND-BY": motivo,
            }
        )
    return lista


def generar_desviaciones():
    desc = [
        "Problema con perno crítico",
        "Falla de energía en área de chancado",
        "Interferencia de contratistas",
        "Bloqueo de acceso por movimiento de equipos",
        "Retraso en llegada de repuestos",
        "Condiciones climáticas adversas (nieve/lluvia)",
    ]
    n = random.randint(1, 4)
    items = random.sample(desc, n)
    lista = []
    for idx, d in enumerate(items, start=1):
        inicio = random.uniform(0.05, 0.25)
        fin = inicio + random.uniform(0.03, 0.15)
        lista.append(
            {
                "N": idx,
                "DESCRIPCIÓN": d,
                "H. INICIO": round(inicio, 5),
                "H. FIN": round(fin, 5),
                "R. CRÍT": "Sí" if random.random() > 0.7 else "No",
            }
        )
    return lista


def generar_actividades():
    base = [
        "Izaje y retiro de componente crítico",
        "Revisión estructural de bastidor",
        "Ajuste y torqueo de pernos",
        "Lubricación de componentes móviles",
        "Limpieza y orden del área de trabajo",
        "Chequeo de holguras y alineamiento",
    ]
    random.shuffle(base)
    acts = base[: random.randint(3, 6)]
    return [{"N": i + 1, "DESCRIPCIÓN": a} for i, a in enumerate(acts)]


def generar_hallazgos():
    base = [
        "Uso incorrecto de EPP en área de izaje",
        "Objetos en pasillos de circulación",
        "Trabajador sin línea de vida en altura",
        "Procedimiento de bloqueo/etiquetado incompleto",
        "Zona caliente sin demarcación",
    ]
    n = random.randint(0, 3)
    items = random.sample(base, n)
    return [{"N": i + 1, "DESCRIPCIÓN": x} for i, x in enumerate(items)]


def generar_hitos_entrante():
    base = [
        "Montaje de componente crítico en próximo turno",
        "Entrega parcial de repuestos críticos",
        "Validación metrológica pendiente",
        "Cierre de inspección y pruebas funcionales",
    ]
    items = random.sample(base, random.randint(1, 3))
    return [{"N": i + 1, "DESCRIPCIÓN": a} for i, a in enumerate(items)]


# =====================================================================
# 6) CONSTRUIR CONTEXTO POR DOCUMENTO DESDE estructura_reporte.json
# =====================================================================
def _get_dropdown_values(estructura, tag_name):
    for c in estructura.get("controles", []):
        if c.get("tag") == tag_name or c.get("alias") == tag_name:
            return c.get("valores") or []
    return []


def build_context_from_estructura(estructura, idx, total_docs):
    proyecto = random.choice(PROYECTOS)
    ref, ubic = random.choice(REFERENCIAS_UBICACIONES)
    start, end = random.choice(TURNOS_POSSIBLES)
    jornada = "Día" if int(start.split(":")[0]) < 12 else "Noche"

    opciones_jornada = _get_dropdown_values(estructura, "tipo_jornada")
    if opciones_jornada:
        if jornada not in opciones_jornada:
            jornada = random.choice(opciones_jornada)
    else:
        opciones_jornada = ["Día", "Noche"]

    base_date = datetime(2025, 1, 1)
    fecha_turno = (base_date + timedelta(days=idx - 1)).strftime("%Y-%m-%d")

    rep_metso, cargo_metso = random.choice(REPRESENTANTES_METSO)
    rep_cli, cargo_cli = random.choice(REPRESENTANTES_CLIENTE)

    context = {
        "nombre_proyecto": proyecto,
        "referencia_turno": ref,
        "ubicacion_turno": ubic,
        "fecha_turno": fecha_turno,
        "num_reporte_diario": str(idx),
        "tipo_jornada": jornada,
        "hora_inicio": start,
        "hora_termino": end,
        "nombre_representante_metso_reporte_diario": rep_metso,
        "cargo_representante_reporte_diario": cargo_metso,
        "fecha_representante_metso_reporte_diario": fecha_turno,
        "nombre_representante_cliente_reporte_diario": rep_cli,
        "cargo_representante_cliente_reporte_diario": cargo_cli,
        "fecha_representante_cliente_reporte_diario": fecha_turno,
        "excel_tables": {
            "personal_directo": generar_personal_directo(),
            "personal_indirecto": generar_personal_indirecto(),
            "equipos_auxiliares": generar_equipos_auxiliares(),
            "desviaciones_turno": generar_desviaciones(),
            "act_realizadas": generar_actividades(),
            "hallazgos_seguridad": generar_hallazgos(),
            "hitos_relevantes_turno": generar_hitos_entrante(),
        },
    }

    return context


# =====================================================================
# 7) GENERAR REPORTES A PARTIR DE PLANTILLA + CONTEXTOS
# =====================================================================
def generate_reports(
    template_path, output_dir, docs_context, daily_plan, daily_real, base_name="REPORTE"
):
    os.makedirs(output_dir, exist_ok=True)

    total_docs = len(docs_context)
    if len(daily_plan) != total_docs or len(daily_real) != total_docs:
        raise ValueError("daily_plan y daily_real deben tener largo = N")

    if abs(sum(daily_plan) - 1.0) > 1e-6:
        raise ValueError("La suma de daily_plan debe ser 1.0")

    if abs(sum(daily_real) - 1.0) > 1e-6:
        raise ValueError("La suma de daily_real debe ser 1.0")

    print(f"Iniciando generación de {total_docs} reportes en '{output_dir}/'")

    with open(template_path, "rb") as f:
        template_bytes = f.read()

    for i, context in enumerate(docs_context, start=1):
        out_path = os.path.join(output_dir, f"{base_name}_{i:02d}.docx")

        cc_context = {k: v for k, v in context.items() if k != "excel_tables"}
        tables_ctx = context.get("excel_tables", {})

        with zipfile.ZipFile(BytesIO(template_bytes), "r") as zin:
            with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)

                    if item.filename == "word/document.xml":
                        data = fill_content_controls(data, cc_context)

                    elif item.filename.startswith(
                        "word/embeddings/"
                    ) and item.filename.endswith(".xlsx"):

                        hh_turno = _calculate_shift_hours(
                            cc_context["hora_inicio"], cc_context["hora_termino"]
                        )

                        name_updates = {
                            "hh_turno_directo": hh_turno / 24.0,
                            "hh_turno_indirecto": hh_turno / 24.0,
                        }
                        data = update_embedded_xlsx_names(data, name_updates)

                        data = update_avance_turno_xlsx(
                            data,
                            doc_index=i,
                            total_docs=total_docs,
                            daily_plan=daily_plan,
                            daily_real=daily_real,
                        )

                        data = update_embedded_xlsx_tables(data, tables_ctx)

                    zout.writestr(item, data)

        try:
            num_rep = int(cc_context.get("num_reporte_diario", i))
        except Exception:
            num_rep = i

        _insertar_evidencias_en_docx(out_path, num_rep, total_evidencias=3)

        print(f"✔ Generado: {out_path}")

    print("\nProceso completado.")


# =====================================================================
# 8) MAIN: USO DIRECTO CON estructura_reporte.json
# =====================================================================
if __name__ == "__main__":
    ESTRUCTURA_JSON = "estructura_reporte.json"
    OUTPUT_DIR = "salida_reportes_desde_json"
    N = 7# Total de turnos/documentos a generar

    with open(ESTRUCTURA_JSON, "r", encoding="utf-8") as f:
        estructura = json.load(f)

    plantilla = estructura.get("plantilla", "REPORTE.docx")
    TEMPLATE = os.path.join(".", plantilla)

    daily_plan, daily_real = generar_curvas_s_clasica(N)

    docs_context = []
    for i in range(1, N + 1):
        ctx = build_context_from_estructura(estructura, i, N)
        docs_context.append(ctx)

    try:
        generate_reports(TEMPLATE, OUTPUT_DIR, docs_context, daily_plan, daily_real)
    except Exception as e:
        print("\n❌ ERROR:", e)
