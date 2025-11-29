
# C:\Users\jonat\Desktop\memoria\Diseño\Formatos Documentos - Final\1 LISTOS\FINAL\Pre\test\Reporte Diario\llenar_reporte_diario.py
import os
import zipfile
from io import BytesIO
import xml.etree.ElementTree as ET
from copy import deepcopy  # Necesario para modificar diccionarios de contexto
import re
from openpyxl import load_workbook
from openpyxl.utils import range_boundaries, get_column_letter
from openpyxl.utils.exceptions import InvalidFileException
import math
import random
from datetime import datetime, timedelta

# --- NUEVO: para generar la imagen de evidencia e insertarla en Word ---
from tempfile import NamedTemporaryFile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml.ns import qn as qn_docx
from docx.shared import Inches

# Espacio de nombres de Word para la búsqueda en el XML (document.xml)
NS_W = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# ... (código anterior permanece igual)

# ... (código anterior permanece igual)

def _insertar_evidencias_en_docx(doc_path, num_reporte_diario, total_evidencias=4):
    """
    Inserta múltiples evidencias fotográficas en el documento.
    Duplica el bloque SDT que contiene:
        - num_imagen_evidencia
        - imagen_evidencia
        - descripcion_imagen_evidencia
    """
    doc = Document(doc_path)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # 1) Buscar el SDT principal que contiene el bloque completo
    # Cambiar la búsqueda para usar el tag del bloque de repetición principal
    bloque_principal = None
    for sdt in doc.element.body.findall(".//w:sdt", ns):
        tags = [t.get(qn_docx("w:val")) for t in sdt.findall(".//w:tag", ns)]
        if "bloque_imagen_evidencia" in tags:  # Buscar por el tag del bloque de repetición
            bloque_principal = sdt
            break

    if bloque_principal is None:
        print("❌ No se encontró el bloque SDT de evidencia con tag 'bloque_imagen_evidencia'.")
        return

    # 2) Guardar copia del bloque original
    from copy import deepcopy
    bloque_template = deepcopy(bloque_principal)

    # 3) Eliminar el bloque original del documento
    parent = bloque_principal.getparent()
    parent.remove(bloque_principal)

    # 4) Insertar N evidencias
    for i in range(1, total_evidencias + 1):
        # Clonar bloque base
        new_block = deepcopy(bloque_template)
        parent.append(new_block)

        # Títulos para imagen generada (dependen de i)
        titulo_imagen = (
            f"Foto Evidencia {i} de {total_evidencias}\n"
            f"Reporte Turno {num_reporte_diario}"
        )
        desc_imagen = (
            f"Registro fotográfico {i} de {total_evidencias} "
            f"correspondiente al Reporte Diario {num_reporte_diario}."
        )

        # Generar imagen específica para esta evidencia (mover dentro del loop)
        img_path = _crear_imagen_evidencia(titulo_imagen)

        # --- Rellenar cada SDT interno ---
        for sdt in new_block.findall(".//w:sdt", ns):
            tag = sdt.find(".//w:tag", ns)
            if tag is None:
                continue

            tag_val = tag.get(qn_docx("w:val"))

            # Número
            if tag_val == "num_imagen_evidencia":
                for t in sdt.findall(".//w:t", ns):
                    t.text = ""
                t = sdt.find(".//w:t", ns)
                if t is not None:
                    t.text = str(i)

            # Descripción
            if tag_val == "descripcion_imagen_evidencia":
                for t in sdt.findall(".//w:t", ns):
                    t.text = ""
                t = sdt.find(".//w:t", ns)
                if t is not None:
                    t.text = desc_imagen

            # Imagen: Usar lógica mejorada para insertar dentro del SDT
            if tag_val == "imagen_evidencia":
                # Eliminar dibujo anterior
                for d in sdt.findall(".//w:drawing", ns):
                    d.getparent().remove(d)
                # Borrar cualquier texto
                for t in sdt.findall(".//w:t", ns):
                    t.text = ""

                # Insertar nueva imagen dentro del SDT
                content = sdt.find(".//w:sdtContent", ns)
                if content is not None:
                    paragraphs = content.findall(".//w:p", ns)
                    if paragraphs:
                        # Usar el primer párrafo existente
                        p = paragraphs[0]
                        # Crear un nuevo párrafo para la imagen (sin añadirlo al doc aún)
                        para_obj = doc._body.add_paragraph()  # Corregido: usar doc._body.add_paragraph()
                        para_obj.alignment = 1
                        run = para_obj.add_run()
                        run.add_picture(img_path, width=Inches(5.5))
                        # Insertar después del párrafo existente
                        p.addnext(para_obj._p)
                    else:
                        # Si no hay párrafos, crear uno nuevo
                        para_obj = doc._body.add_paragraph()  # Corregido: usar doc._body.add_paragraph()
                        para_obj.alignment = 1
                        run = para_obj.add_run()
                        run.add_picture(img_path, width=Inches(5.5))
                        content.append(para_obj._p)

    # Guardar documento final
    doc.save(doc_path)

# ... (el resto del código permanece igual)



# =====================================================================
# AUXILIAR: CALCULAR HORAS DE TURNO
# =====================================================================
def _calculate_shift_hours(start_time_str, end_time_str):
    """
    Calcula la duración del turno en horas decimales (hh:mm -> float).
    Maneja turnos que cruzan la medianoche.
    Ejemplo: 21:00 a 09:00 = 12 horas.
    """
    try:
        start_time = datetime.strptime(start_time_str, "%H:%M")
        end_time = datetime.strptime(end_time_str, "%H:%M")

        if end_time < start_time:
            end_time += timedelta(days=1)

        duration = end_time - start_time
        total_hours = duration.total_seconds() / 3600.0
        return total_hours
    except ValueError as e:
        raise ValueError(
            f"Formato de hora inválido ('{start_time_str}' o '{end_time_str}'). "
            f"Debe ser HH:MM. Error: {e}"
        )


# =====================================================================
# ACTUALIZAR NOMBRES DEFINIDOS EN XLSX (HH_TURNO)
# =====================================================================
def update_embedded_xlsx_names(xlsx_bytes, name_values):
    """
    Actualiza los valores de los nombres definidos de Excel.
    name_values es un diccionario {nombre: valor}
    """
    try:
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=False)
    except InvalidFileException:
        return xlsx_bytes

    modified = False

    for name, name_obj in wb.defined_names.items():
        if name in name_values:
            try:
                dest_sheet, dest_coord = next(name_obj.destinations)
            except StopIteration:
                continue

            ws = wb[dest_sheet]
            cell_ref = dest_coord
            cell = ws[cell_ref]
            cell.value = name_values[name]
            cell.number_format = "[h]:mm"
            modified = True

    if modified:
        out = BytesIO()
        wb.save(out)
        return out.getvalue()

    return xlsx_bytes


# =====================================================================
# AJUSTE DE FÓRMULAS (cuando se insertan filas en tablas)
# =====================================================================
def _shift_formula_rows(formula, src_row, dst_row):
    """
    Ajusta referencias A1 en fórmulas.
    Las referencias estructuradas de tabla ([[#This Row],...]) NO deben tocarse.
    """
    if "[#This Row]" in formula or "[[" in formula:
        return formula

    row_offset = dst_row - src_row
    pattern = re.compile(r"(\$?[A-Z]{1,3})(\d+)")

    def repl(match):
        col = match.group(1)
        row = int(match.group(2))
        new_row = row + row_offset
        return f"{col}{new_row}"

    return pattern.sub(repl, formula)


# =====================================================================
# PEQUEÑAS UTILIDADES
# =====================================================================
def _ref_to_bounds(ref):
    return range_boundaries(ref)


def _get_table_headers(ws, ref):
    min_col, min_row, max_col, max_row = _ref_to_bounds(ref)
    header_row = min_row
    headers = {}

    for col in range(min_col, max_col + 1):
        cell = ws.cell(row=header_row, column=col)
        if cell.value:
            headers[str(cell.value).strip()] = col

    return headers, min_row, max_row


def _is_hour_format_valid(value):
    try:
        f_val = float(value)
        return f_val >= 0 and f_val < 1
    except (ValueError, TypeError):
        return False


def _is_integer_valid(value):
    try:
        i_val = int(value)
        return i_val >= 0
    except (ValueError, TypeError):
        return False


# =====================================================================
# EXPANDIR TABLA EXCEL AUTOMÁTICAMENTE
# =====================================================================
def ensure_table_min_rows(ws, tbl, required_rows):
    """
    Garantiza que la tabla Excel tenga al menos 'required_rows' filas de datos.
    Si no, inserta filas, actualiza tbl.ref y copia formato/fórmulas
    de la PRIMERA FILA DE DATOS de la tabla.
    """
    min_col, min_row, max_col, max_row = range_boundaries(tbl.ref)
    header_row = min_row
    first_data_row = header_row + 1

    current_rows = max_row - header_row
    missing = required_rows - current_rows
    if missing <= 0:
        return header_row, max_row

    insert_at = max_row + 1
    ws.insert_rows(insert_at, missing)
    new_max_row = max_row + missing

    for col in range(min_col, max_col + 1):
        src_cell = ws.cell(row=first_data_row, column=col)
        for new_row in range(insert_at, new_max_row + 1):
            dst_cell = ws.cell(row=new_row, column=col)
            dst_cell._style = src_cell._style

            src_val = src_cell.value
            if isinstance(src_val, str) and src_val.startswith("="):
                if "[#This Row]" in src_val or "[[" in src_val:
                    dst_cell.value = src_val
                else:
                    dst_cell.value = _shift_formula_rows(
                        src_val, first_data_row, new_row
                    )
            else:
                dst_cell.value = None

    new_ref = (
        f"{get_column_letter(min_col)}{min_row}:"
        f"{get_column_letter(max_col)}{new_max_row}"
    )
    tbl.ref = new_ref
    return header_row, new_max_row


# =====================================================================
# 1) RELLENAR CONTROLES DE CONTENIDO (document.xml) POR TAG
# =====================================================================
def fill_content_controls(document_xml_bytes, context):
    root = ET.fromstring(document_xml_bytes)

    for sdt in root.findall(".//w:sdt", NS_W):
        tag_elem = sdt.find("./w:sdtPr/w:tag", NS_W)
        if tag_elem is None:
            continue

        tag_val = tag_elem.get(f"{{{NS_W['w']}}}val")

        if (
            tag_val == "nombre_representante_reporte_diario"
            and "nombre_representante_metso_reporte_diario" in context
        ):
            tag_val = "nombre_representante_metso_reporte_diario"

        if not tag_val or tag_val not in context:
            continue

        new_text = str(context[tag_val])
        texts = sdt.findall(".//w:sdtContent//w:t", NS_W)
        if not texts:
            continue

        for t in texts:
            t.text = ""
        texts[0].text = new_text

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


# =====================================================================
# 2) ACTUALIZAR TABLA 'avance_turno' (curva S) – NO TOCAR LÓGICA
# =====================================================================
def update_avance_turno_xlsx(xlsx_bytes, doc_index, total_docs, daily_plan, daily_real):
    col_name_plan = "% AVANCE DIARIO PLANIFICADO"
    col_name_real = "% AVANCE DIARIO REAL"
    col_name_acum_plan = "AVANCE PLANIFICADO (ACUMULADO)"
    col_name_acum_real = "AVANCE REAL (ACUMULADO)"
    col_name_n_reporte = "N° REPORTE DIARIO"

    try:
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=False)
    except InvalidFileException:
        return xlsx_bytes

    found = False

    for ws in wb.worksheets:
        for tbl in ws._tables.values():
            if tbl.name != "avance_turno":
                continue

            found = True
            header_row, max_row = ensure_table_min_rows(ws, tbl, total_docs)

            ref = tbl.ref
            headers, header_row, max_row = _get_table_headers(ws, ref)
            total_rows = max_row - header_row

            col_map = {
                col_name_n_reporte: headers.get(col_name_n_reporte),
                col_name_plan: headers.get(col_name_plan),
                col_name_real: headers.get(col_name_real),
                col_name_acum_plan: headers.get(col_name_acum_plan),
                col_name_acum_real: headers.get(col_name_acum_real),
            }

            if not all(col_map.values()):
                return xlsx_bytes

            for i in range(1, total_rows + 1):
                excel_row = header_row + i

                if i <= doc_index:
                    p = float(daily_plan[i - 1])
                    r = float(daily_real[i - 1])
                    acum_p = sum(daily_plan[:i])
                    acum_r = sum(daily_real[:i])

                    ws.cell(excel_row, col_map[col_name_n_reporte]).value = i
                    ws.cell(excel_row, col_map[col_name_plan]).value = p
                    ws.cell(excel_row, col_map[col_name_real]).value = r

                    if i == total_docs:
                        ws.cell(excel_row, col_map[col_name_acum_plan]).value = 1.0
                        ws.cell(excel_row, col_map[col_name_acum_real]).value = 1.0
                    else:
                        ws.cell(excel_row, col_map[col_name_acum_plan]).value = acum_p
                        ws.cell(excel_row, col_map[col_name_acum_real]).value = acum_r

                elif i <= total_docs:
                    for col in col_map.values():
                        ws.cell(excel_row, col).value = None
                else:
                    for col in col_map.values():
                        ws.cell(excel_row, col).value = None

    if not found:
        return xlsx_bytes

    out = BytesIO()
    wb.save(out)
    return out.getvalue()


# =====================================================================
# 3) ACTUALIZAR OTRAS TABLAS EMBEBIDAS
# =====================================================================
def update_data_rows(ws, table_name, data_rows, header_row, max_row, headers):
    HOUR_COLS = ["H.M. ASIGN.", "H. M. OPER.", "H.M. S-BY", "H. INICIO", "H. FIN"]
    INT_COLS = ["CANT", "CANT. PLAN.", "CANT. REAL"]

    for idx, row_data in enumerate(data_rows):
        excel_row = header_row + idx + 1
        if excel_row > max_row:
            continue

        for col_name, value in row_data.items():
            col_name_strip = str(col_name).strip()
            if col_name_strip not in headers:
                continue

            col = headers[col_name_strip]

            if col_name_strip in HOUR_COLS:
                if not _is_hour_format_valid(value):
                    raise ValueError(
                        f"Tabla '{table_name}': '{col_name_strip}' fila {idx+1} debe "
                        f"ser float >=0 <1. Valor: {value}"
                    )

            if col_name_strip in INT_COLS:
                if not _is_integer_valid(value):
                    raise ValueError(
                        f"Tabla '{table_name}': '{col_name_strip}' fila {idx+1} debe "
                        f"ser entero >=0. Valor: {value}"
                    )

            if col_name_strip in [
                "HH PLAN.",
                "HH REAL",
                "H. TOTAL",
                "DISPONIBILIDAD",
                "UTILIZACIÓN",
                "USO EFECTIVO",
                "STAND-BY",
            ]:
                continue

            ws.cell(excel_row, col).value = value

    rows_to_clear_start = header_row + len(data_rows) + 1
    for row in range(rows_to_clear_start, max_row + 1):
        for col in headers.values():
            ws.cell(row=row, column=col).value = None


def update_embedded_xlsx_tables(xlsx_bytes, tables_context):
    try:
        wb = load_workbook(BytesIO(xlsx_bytes), data_only=False)
    except InvalidFileException:
        return xlsx_bytes

    modified = False

    for ws in wb.worksheets:
        for tbl in ws._tables.values():
            name = tbl.name
            if name not in tables_context:
                continue
            if name == "avance_turno":
                continue

            modified = True
            rows = tables_context[name]
            required_rows = len(rows)

            try:
                if required_rows > 0:
                    ensure_table_min_rows(ws, tbl, required_rows)

                ref = tbl.ref
                headers, header_row, max_row = _get_table_headers(ws, ref)
                update_data_rows(ws, name, rows, header_row, max_row, headers)
            except ValueError as e:
                raise e
            except Exception as e:
                print(f"Error procesando tabla '{name}': {e}")

    if modified:
        out = BytesIO()
        wb.save(out)
        return out.getvalue()

    return xlsx_bytes


# =====================================================================
# 4A) GENERAR IMAGEN DE EVIDENCIA
# =====================================================================
def _crear_imagen_evidencia(texto):
    """
    Genera una imagen PNG de fondo blanco con texto centrado en negro.
    Devuelve la ruta del archivo temporal.
    """
    from PIL import Image, ImageDraw, ImageFont
    from tempfile import NamedTemporaryFile

    # Tamaño de la imagen
    width, height = 1200, 800
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    # Fuente
    try:
        font = ImageFont.truetype("arial.ttf", 48)
    except Exception:
        font = ImageFont.load_default()

    # Procesar líneas
    lines = texto.split("\n")
    line_sizes = []

    # Medir cada línea usando textbbox()
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        line_sizes.append((w, h))

    total_h = sum(h for _, h in line_sizes) + (len(lines) - 1) * 10
    y = (height - total_h) // 2

    # Dibujar texto centrado
    for (w, h), line in zip(line_sizes, lines):
        x = (width - w) // 2
        draw.text((x, y), line, fill="black", font=font)
        y += h + 10

    # Guardar en archivo temporal
    tmp = NamedTemporaryFile(delete=False, suffix=".png")
    img.save(tmp.name, "PNG")
    return tmp.name


# =====================================================================
# 4A-bis) REEMPLAZAR DIBUJO DUMMY DE UN SDT POR UNA IMAGEN (tipo curva S)
# =====================================================================
def _replace_sdt_with_image(doc, tag_name, image_path, width_inches=5.5):
    """
    Reemplaza el contenido gráfico (drawing) de un SDT existente identificado
    por su w:tag/@w:val == tag_name, insertando la imagen en el mismo marco.

    Misma lógica que en llenar_informe_tecnico.py (Curva S).
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Iterar todos los controles de contenido
    for sdt in doc.element.body.iter(qn_docx("w:sdt")):
        tag = sdt.find(".//w:tag", ns)
        if tag is None:
            continue

        if tag.get(qn_docx("w:val")) != tag_name:
            continue

        # 1) Borrar cualquier texto dentro del SDT
        for t in sdt.findall(".//w:t", ns):
            t.text = ""

        # 2) Borrar cualquier dibujo (imagen dummy) existente
        for drawing in sdt.findall(".//w:drawing", ns):
            parent = drawing.getparent()
            parent.remove(drawing)

        # 3) Ubicar w:sdtContent
        content = sdt.find(".//w:sdtContent", ns)
        if content is None:
            return

        # 4) Buscar un párrafo existente dentro del sdtContent
        paragraphs = content.findall(".//w:p", ns)

        if paragraphs:
            # Tomamos el primer párrafo del recuadro
            p = paragraphs[0]

            # Creamos un párrafo real de python-docx para la imagen
            # usando el mismo "parent" de los párrafos del documento
            para_obj = doc.paragraphs[0]._parent.add_paragraph()
            para_obj.alignment = 1
            run = para_obj.add_run()
            run.add_picture(image_path, width=Inches(width_inches))

            # Insertamos este párrafo inmediatamente después del p del SDT
            p.addnext(para_obj._p)
        else:
            # Si por alguna razón no hay párrafos, creamos uno y lo
            # colgamos de sdtContent
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            content.append(para._p)

        return  # Solo reemplaza el primer SDT que coincida


# =====================================================================
# 4A) INSERTAR EVIDENCIA EN EL DOCX (USANDO EL SDT EXISTENTE)
# =====================================================================
def _insertar_evidencia_en_docx(doc_path, num_reporte_diario):
    """
    Abre el DOCX generado, crea una imagen de evidencia y rellena:
    - num_imagen_evidencia
    - imagen_evidencia (insertando la imagen EN el SDT existente)
    - descripcion_imagen_evidencia
    usando los controles de contenido existentes.
    """
    doc = Document(doc_path)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Definimos cuántas evidencias "conceptuales" tendrá este reporte
    total_evidencias = random.randint(4, 6)
    num_imagen = 1  # por ahora llenamos solo la evidencia 1 en la plantilla

    titulo_imagen = (
        f"Foto Evidencia {num_imagen} de {total_evidencias}\n"
        f"Reporte Turno {num_reporte_diario}"
    )
    desc_imagen = (
        f"Registro fotográfico {num_imagen} de {total_evidencias} "
        f"correspondiente al Reporte Diario {num_reporte_diario}."
    )

    img_path = _crear_imagen_evidencia(titulo_imagen)

    # 1) Rellenar los SDT de texto (número de imagen y descripción)
    for sdt in doc.element.body.iter(qn_docx("w:sdt")):
        tag_elem = sdt.find(".//w:tag", ns)
        if tag_elem is None:
            continue

        tag_val = tag_elem.get(qn_docx("w:val"))

        # 1) NUMERO DE IMAGEN
        if tag_val == "num_imagen_evidencia":
            texts = sdt.findall(".//w:t", ns)
            for t in texts:
                t.text = ""
            if texts:
                texts[0].text = str(num_imagen)

        # 2) DESCRIPCION DE LA IMAGEN
        elif tag_val == "descripcion_imagen_evidencia":
            texts = sdt.findall(".//w:t", ns)
            for t in texts:
                t.text = ""
            if texts:
                texts[0].text = desc_imagen

    # 2) Reemplazar el dibujo dummy del SDT imagen_evidencia por la imagen creada
    _replace_sdt_with_image(doc, "imagen_evidencia", img_path, width_inches=5.5)

    doc.save(doc_path)


# =====================================================================
# 4B) GENERAR N DOCUMENTOS (REPORTE DIARIO)
# =====================================================================
def generate_reports(
    template_path, output_dir, docs_context, daily_plan, daily_real, base_name="REPORTE"
):
    os.makedirs(output_dir, exist_ok=True)

    total_docs = len(docs_context)
    if len(daily_plan) != total_docs or len(daily_real) != total_docs:
        raise ValueError("daily_plan y daily_real deben tener largo = N")

    if abs(sum(daily_plan) - 1.0) > 1e-6:
        raise ValueError("La suma de daily_plan debe ser 1.0")

    if abs(sum(daily_real) - 1.0) > 1e-6:
        raise ValueError("La suma de daily_real debe ser 1.0")

    print(f"Iniciando generación de {total_docs} reportes en '{output_dir}/'")

    with open(template_path, "rb") as f:
        template_bytes = f.read()

    for i, context in enumerate(docs_context, start=1):
        out_path = os.path.join(output_dir, f"{base_name}_{i:02d}.docx")

        cc_context = {k: v for k, v in context.items() if k != "excel_tables"}
        tables_ctx = context.get("excel_tables", {})

        with zipfile.ZipFile(BytesIO(template_bytes), "r") as zin:
            with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)

                    if item.filename == "word/document.xml":
                        data = fill_content_controls(data, cc_context)

                    elif item.filename.startswith(
                        "word/embeddings/"
                    ) and item.filename.endswith(".xlsx"):

                        # 1) Nombres definidos HH_TURNO
                        hh_turno = _calculate_shift_hours(
                            cc_context["hora_inicio"], cc_context["hora_termino"]
                        )

                        name_updates = {
                            "hh_turno_directo": hh_turno / 24.0,
                            "hh_turno_indirecto": hh_turno / 24.0,
                        }
                        data = update_embedded_xlsx_names(data, name_updates)

                        # 2) Curva S / Avance turno
                        data = update_avance_turno_xlsx(
                            data,
                            doc_index=i,
                            total_docs=total_docs,
                            daily_plan=daily_plan,
                            daily_real=daily_real,
                        )

                        # 3) Otras tablas
                        data = update_embedded_xlsx_tables(data, tables_ctx)

                    zout.writestr(item, data)

        # --- NUEVO: insertar registro fotográfico de evidencia ---
        try:
            num_rep = int(cc_context.get("num_reporte_diario", i))
        except Exception:
            num_rep = i

        _insertar_evidencias_en_docx(out_path, num_rep, total_evidencias=3)

        print(f"✔ Generado: {out_path}")

    print("\nProceso completado.")


# =====================================================================
# 4C) CURVA S CLÁSICA (para daily_plan / daily_real)
# =====================================================================
def generar_curvas_s_clasica(N, ruido_real=0.10):
    """
    Genera daily_plan y daily_real con forma de curva S (sigmoidal).
    - daily_plan: suave, simétrica, suma 1.0
    - daily_real: daily_plan + ruido, normalizado a 1.0
    """
    k = N / 2
    a = N / 10

    S = [1 / (1 + math.exp(-(i - k) / a)) for i in range(N + 1)]
    S_min = S[0]
    S_max = S[-1]
    S = [(x - S_min) / (S_max - S_min) for x in S]

    daily_plan = [S[i] - S[i - 1] for i in range(1, N + 1)]
    total = sum(daily_plan)
    daily_plan = [x / total for x in daily_plan]

    daily_real = []
    for p in daily_plan:
        delta = random.uniform(-ruido_real, ruido_real)
        daily_real.append(max(0.0, p + delta))

    total_r = sum(daily_real)
    daily_real = [r / total_r for r in daily_real]

    return daily_plan, daily_real

# =====================================================================
# 5) EJEMPLO DE USO CON DATOS DINÁMICOS
# =====================================================================
if __name__ == "__main__":
    TEMPLATE = "REPORTE.docx"
    OUTPUT_DIR = "salida_reportes_completo"
    N = 6  # Total de turnos/documentos

    daily_plan, daily_real = generar_curvas_s_clasica(N)
    docs_context = []

    shifts = [
        ("07:00", "19:00"),
        ("08:00", "20:00"),
        ("06:30", "18:30"),
        ("21:00", "09:00"),
        ("22:00", "10:00"),
        ("20:30", "08:30"),
        ("07:30", "19:30"),
    ]

    referencias_ubicaciones = [
        ("CH-001", "CHANCADORA PRIMARIA - ZONA NORTE"),
        ("CH-002", "CHANCADORA SECUNDARIA - ZONA SUR"),
        ("SG-001", "MOLINO SAG - ZONA ESTE"),
    ]

    def generar_personal_directo(turno):
        base = [
            ("Supervisor Mecánico", 1),
            ("Técnico Mecánico", 4),
            ("Soldador Especialista", 2),
            ("Ayudante Mecánico", 3),
            ("Lubricador", 1),
            ("Técnico Hidráulico", 1),
        ]
        personal = []
        for cargo, base_cant in base:
            variacion = random.choice([-1, 0, 1])
            cant_plan = max(1, base_cant + variacion)
            cant_real = max(0, cant_plan + random.choice([-1, 0, 1]))
            personal.append(
                {
                    "CANT": cant_plan,
                    "CARGO": cargo,
                    "CANT. PLAN.": cant_plan,
                    "CANT. REAL": cant_real,
                }
            )
        return personal

    def generar_personal_indirecto(turno):
        cargos = [
            "HSE",
            "Bodeguero",
            "Planificador",
            "Ingeniero de Proyecto",
            "Administrador de Contrato",
        ]
        personal = []
        for cargo in cargos:
            cant_plan = 1
            cant_real = 1 if random.random() > 0.2 else 0
            personal.append(
                {
                    "CANT": cant_plan,
                    "CARGO": cargo,
                    "CANT. PLAN.": cant_plan,
                    "CANT. REAL": cant_real,
                }
            )
        return personal

    def generar_equipos_auxiliares(turno):
        equipos = [
            ("Camión Pluma", 0.45, 0.50),
            ("Camioneta", 0.50, 0.50),
            ("Grúa Horquilla", 0.40, 0.48),
            ("Plataforma Elevadora", 0.35, 0.45),
            ("Compresor", 0.50, 0.50),
            ("Torre de Iluminación", 0.45, 0.50),
        ]
        motivos = [
            "Condición de terreno",
            "Esperando instrucciones",
            "Clima adverso",
            "Interferencia de otros equipos",
            "Falta de operador",
        ]
        lista = []
        for nombre, min_op, max_op in equipos:
            asign = 0.50
            oper = round(random.uniform(min_op, max_op), 3)
            standby = round(max(0.0, asign - oper), 3)
            motivo = random.choice(motivos) if standby > 0 else ""
            lista.append(
                {
                    "CANT": random.choice([1, 1, 2]),
                    "DESCRIPCIÓN": nombre,
                    "H.M. ASIGN.": asign,
                    "H. M. OPER.": oper,
                    "H.M. S-BY": standby,
                    "OBS. / MOTIVO STAND-BY": motivo,
                }
            )
        return lista

    def generar_desviaciones(turno):
        desc = [
            "Problema con perno",
            "Falla de energía",
            "Interferencia de contratistas",
            "Bloqueo de acceso",
            "Retraso en repuestos",
            "Condiciones climáticas adversas",
        ]
        n = random.randint(1, 4)
        items = random.sample(desc, n)
        lista = []
        for idx, d in enumerate(items, start=1):
            inicio = random.uniform(0.05, 0.25)
            fin = inicio + random.uniform(0.03, 0.15)
            lista.append(
                {
                    "N": idx,
                    "DESCRIPCIÓN": d,
                    "H. INICIO": round(inicio, 5),
                    "H. FIN": round(fin, 5),
                    "R. CRÍT": "Sí" if random.random() > 0.7 else "No",
                }
            )
        return lista

    def generar_actividades(turno):
        base = [
            "Izaje y retiro de pieza",
            "Revisión estructural",
            "Ajuste de pernos",
            "Lubricación de componentes",
            "Limpieza y orden",
            "Chequeo de torque",
        ]
        random.shuffle(base)
        acts = base[: random.randint(3, 6)]
        return [{"N": i + 1, "DESCRIPCIÓN": a} for i, a in enumerate(acts)]

    def generar_hallazgos(turno):
        base = [
            "Uso incorrecto de EPP",
            "Objetos en pasillos",
            "Trabajador sin amarre",
            "Procedimiento no seguido",
            "Zona sin demarcar",
        ]
        n = random.randint(0, 3)
        items = random.sample(base, n)
        return [{"N": i + 1, "DESCRIPCIÓN": x} for i, x in enumerate(items)]

    def generar_hitos(turno):
        base = [
            "Montaje de equipo crítico",
            "Entrega parcial de repuestos",
            "Validación metrológica",
            "Cierre de inspección",
        ]
        items = random.sample(base, random.randint(1, 3))
        return [{"N": i + 1, "DESCRIPCIÓN": a} for i, a in enumerate(items)]

    for i in range(1, N + 1):
        start, end = random.choice(shifts)
        jornada = "Día" if int(start.split(":")[0]) < 12 else "Noche"
        ref, ubic = random.choice(referencias_ubicaciones)

        context = {
            "nombre_proyecto": "OVERHAUL CHANCADORA X",
            "referencia_turno": ref,
            "fecha_turno": f"2025-01-0{i}",
            "ubicacion_turno": ubic,
            "num_reporte_diario": str(i),
            "tipo_jornada": jornada,
            "hora_inicio": start,
            "hora_termino": end,
            "nombre_representante_metso_reporte_diario": f"Metso Rep {i}",
            "nombre_representante_cliente_reporte_diario": f"Cliente Rep {i}",
            "excel_tables": {
                "personal_directo": generar_personal_directo(i),
                "personal_indirecto": generar_personal_indirecto(i),
                "equipos_auxiliares": generar_equipos_auxiliares(i),
                "desviaciones_turno": generar_desviaciones(i),
                "act_realizadas": generar_actividades(i),
                "hallazgos_seguridad": generar_hallazgos(i),
                "hitos_relevantes_turno": generar_hitos(i),
            },
        }
        docs_context.append(context)

    try:
        generate_reports(TEMPLATE, OUTPUT_DIR, docs_context, daily_plan, daily_real)
    except Exception as e:
        print("\n❌ ERROR:", e)
